"""Read-only Legal DOCX output for latest human-confirmed KHMT packages."""

from __future__ import annotations

import re
from collections.abc import Iterable, Mapping
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from .candidate_review import CandidateReviewService, ReviewedCandidate
from .khmt_contract import PlanPackage

LEGAL_FIELD_HEADERS = (
    "Mã kế hoạch",
    "Gói tin",
    "Tên gói",
    "Chủ đầu tư",
    "Quyết định phê duyệt",
    "Giá gói thầu",
    "Nguồn vốn",
    "Hình thức lựa chọn",
    "Qua mạng",
    "Sơ tuyển",
    "Phương thức",
    "Hình thức hợp đồng",
    "Thời gian lựa chọn",
    "Thời gian thực hiện",
    "Địa bàn",
)

_RAW_FIELD_BY_LABEL = {
    "Gói tin": "GÓI TIN",
    "Tên gói": "TÊN GÓI THẦU",
    "Chủ đầu tư": "TÊN CHỦ ĐẦU TƯ",
    "Quyết định phê duyệt": "NỘI DUNG PHÊ DUYỆT",
    "Giá gói thầu": "GIÁ GÓI THẦU",
    "Nguồn vốn": "NGUỒN VỐN",
    "Hình thức lựa chọn": "HÌNH THỨC LỰA CHỌN",
    "Qua mạng": "QUA MẠNG",
    "Sơ tuyển": "SƠ TUYỂN",
    "Phương thức": "PHƯƠNG THỨC",
    "Hình thức hợp đồng": "HÌNH THỨC HỢP ĐỒNG",
    "Thời gian lựa chọn": "THỜI GIAN LỰA CHỌN",
    "Thời gian thực hiện": "THỜI GIAN THỰC HIỆN",
    "Địa bàn": "ĐỊA BÀN",
}
_SAFE_PLAN_BASE = re.compile(r"^[^\\/:*?\"<>|\x00-\x1f]+$")


class LegalDocxExportError(ValueError):
    """The confirmed package cannot be rendered safely as a Legal DOCX."""


class LegalDocxCollisionError(LegalDocxExportError):
    """A deterministic DOCX target already exists or is duplicated."""


@dataclass(frozen=True, slots=True)
class LegalDocxExportResult:
    output: Path
    plan_base_id: str
    plan_revision: str | None
    source_row: int


def export_confirmed_legal_docx(
    review_service: CandidateReviewService,
    packages: Iterable[PlanPackage],
    *,
    output_dir: Path = Path("."),
) -> tuple[LegalDocxExportResult, ...]:
    """Write one DOCX per latest confirmed package without mutating MI state."""

    confirmed = review_service.current_confirmed(packages)
    return export_confirmed_legal_docx_records(
        tuple((reviewed.package, reviewed.event) for reviewed in confirmed),
        output_dir=output_dir,
    )


def export_confirmed_legal_docx_records(
    records: Iterable[tuple[PlanPackage, Any]],
    *,
    output_dir: Path = Path("."),
) -> tuple[LegalDocxExportResult, ...]:
    """Render already-selected confirmed KHMT packages without another review authority."""

    selected = tuple(sorted(records, key=lambda pair: _package_sort_key(pair[0])))
    if not selected:
        return ()

    output_dir = Path(output_dir)
    targets = tuple(_target_path_for_package(output_dir, package) for package, _event in selected)
    if len(set(targets)) != len(targets):
        raise LegalDocxCollisionError(
            "Nhiều gói xác nhận cùng mã kế hoạch tạo ra cùng tên DOCX; không ghi đè."
        )
    existing = next((path for path in targets if path.exists()), None)
    if existing is not None:
        raise LegalDocxCollisionError(f"File DOCX đã tồn tại, không ghi đè: {existing.name}")

    output_dir.mkdir(parents=True, exist_ok=True)
    results: list[LegalDocxExportResult] = []
    for (package, _event), target in zip(selected, targets, strict=True):
        _write_new_document(target, package)
        results.append(
            LegalDocxExportResult(
                output=target,
                plan_base_id=package.plan.plan_base_id,
                plan_revision=package.plan.plan_revision,
                source_row=package.source_row,
            )
        )
    return tuple(results)


def _target_path(output_dir: Path, reviewed: ReviewedCandidate) -> Path:
    return _target_path_for_package(output_dir, reviewed.package)


def _target_path_for_package(output_dir: Path, package: PlanPackage) -> Path:
    base_id = package.plan.plan_base_id
    if not base_id or not _SAFE_PLAN_BASE.fullmatch(base_id):
        raise LegalDocxExportError("Mã kế hoạch không hợp lệ để tạo tên file DOCX.")
    return output_dir / f"ThongTin_{base_id}.docx"


def _write_new_document(path: Path, package: PlanPackage) -> None:
    document = Document()
    document.add_heading("Thông tin gói thầu", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in _field_values(package):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    payload = BytesIO()
    document.save(payload)
    try:
        with path.open("xb") as stream:
            stream.write(payload.getvalue())
    except FileExistsError as exc:
        raise LegalDocxCollisionError(f"File DOCX đã tồn tại, không ghi đè: {path.name}") from exc


def _field_values(package: PlanPackage) -> tuple[tuple[str, str], ...]:
    raw_fields: Mapping[str, Any] = package.raw_fields
    values: list[tuple[str, str]] = [("Mã kế hoạch", package.plan.plan_id_raw)]
    for label in LEGAL_FIELD_HEADERS[1:]:
        key = _RAW_FIELD_BY_LABEL[label]
        value = raw_fields.get(key)
        if label == "Địa bàn" and value is None:
            value = package.location_detail_raw
        values.append((label, _display_value(value)))
    return tuple(values)


def _display_value(value: Any) -> str:
    return "" if value is None else str(value)


def _sort_key(reviewed: ReviewedCandidate) -> tuple[str, str, int, str, int]:
    event = reviewed.event
    return (
        event.source_sha256.casefold(),
        event.source_sheet.casefold(),
        event.source_row,
        event.plan_id_raw.casefold(),
        event.id,
    )


def _package_sort_key(package: PlanPackage) -> tuple[str, str, int, str]:
    batch = package.plan.import_batch
    return (
        batch.source_sha256.casefold(),
        batch.sheet.casefold(),
        package.source_row,
        package.plan.plan_id_raw.casefold(),
    )
